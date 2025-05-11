# ✅ STREAMLIT FRONTEND (app.py)
import streamlit as st
import requests
import os

st.set_page_config(page_title="Compliance Form Assistant")
st.title("📄 Compliance Form Assistant")

def log(msg):
    st.text(f"🪵 {msg}")

# ✅ Secrets
APIFY_TOKEN = os.environ["APIFY_TOKEN"]
APIFY_ACTOR_ID = os.environ["APIFY_ACTOR_ID"]  # e.g. "username~actor-name"

uploaded = st.file_uploader("Upload DOCX file", type=["docx"])

if uploaded:
    with open("temp.docx", "wb") as f:
        f.write(uploaded.getbuffer())
    log("✅ File saved locally")

    # Create KV store
    res = requests.post(
        f"https://api.apify.com/v2/key-value-stores?token={APIFY_TOKEN}",
        json={"name": "form-upload-store"}
    )
    store_id = res.json()["data"]["id"]
    log(f"🗂 KV store created: {store_id}")

    # Upload file to INPUT_FILE
    with open("temp.docx", "rb") as f:
        put = requests.put(
            f"https://api.apify.com/v2/key-value-stores/{store_id}/records/INPUT_FILE?token={APIFY_TOKEN}",
            files={"value": ("input.docx", f)}
        )
    if put.status_code not in [200, 201]:
        st.error("❌ Upload failed")
        st.stop()

    # Run actor with inputFileStoreId
    run = requests.post(
        f"https://api.apify.com/v2/acts/{APIFY_ACTOR_ID}/runs?token={APIFY_TOKEN}",
        json={"input": {"inputFileStoreId": store_id}},
        headers={"Content-Type": "application/json"}
    )
    run_id = run.json()["data"]["id"]

    # Poll
    status = "RUNNING"
    while status in ["RUNNING", "READY"]:
        status = requests.get(
            f"https://api.apify.com/v2/actor-runs/{run_id}?token={APIFY_TOKEN}"
        ).json()["data"]["status"]
        log(f"⏳ Status: {status}")

    # Get output
    out = requests.get(
        f"https://api.apify.com/v2/key-value-stores/{store_id}/records/output.json?token={APIFY_TOKEN}"
    ).json()

    st.subheader("🤖 Unanswered Questions")
    for q in out.get("unknown_questions", {}):
        st.text_input(q)

    st.subheader("📥 Download")
    file_name = out["filled_file"].split("/")[-1]
    download = f"https://api.apify.com/v2/key-value-stores/{store_id}/records/{file_name}?token={APIFY_TOKEN}"
    st.markdown(f"[Download Filled Form]({download})")


# ✅ ACTOR BACKEND (main.py)
import os
from apify_client import ApifyClient
from docx import Document

KNOWN_ANSWERS = {
    "Company Name": "BGR, Inc.",
    "Company Address": "6392 Gano Road, West Chester, OH 45069",
    "Company Contact": "Mike Willging",
    "Contact Email": "customerservice@packbgr.com",
    "Job Title": "Warehouse Quality Manager",
    "DUNS": "063974430",
    "Tax ID": "31-0841900"
}

def extract(path):
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip().endswith('?')]

def fill(input_path, output_path, answers):
    doc = Document(input_path)
    for p in doc.paragraphs:
        for k, v in answers.items():
            if k in p.text:
                p.text = p.text.replace(k, f"{k} {v}")
    doc.save(output_path)

def main():
    store_id = os.environ["APIFY_INPUT_JSON"]
    store_id = eval(store_id)["inputFileStoreId"]

    client = ApifyClient(os.environ["APIFY_TOKEN"])
    store = client.key_value_store(store_id)

    file_path = "/tmp/input.docx"
    store.download_record("INPUT_FILE", local_path=file_path)

    q = extract(file_path)
    known, unknown = {}, {}
    for question in q:
        for k, v in KNOWN_ANSWERS.items():
            if k.lower() in question.lower():
                known[question] = v
                break
        else:
            unknown[question] = ""

    out_file = file_path.replace(".docx", "_filled.docx")
    fill(file_path, out_file, known)

    store.upload_file("filled.docx", out_file)
    store.set_record("output", {"filled_file": "filled.docx", "unknown_questions": unknown})

if __name__ == "__main__":
    main()
